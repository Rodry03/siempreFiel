import io
import logging
import os
import tempfile
from datetime import date

from docx import Document

logger = logging.getLogger(__name__)

_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "contracts", "contrato_preadopcion.docx")
_MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

_DATA_FONT = "Times New Roman"


def _apply_font(run):
    run.font.name = _DATA_FONT
    run.font.bold = True


def _set_run(para, run_idx: int, value: str):
    value = value.upper()
    runs = para.runs
    if run_idx < len(runs):
        runs[run_idx].text = value
        _apply_font(runs[run_idx])
    else:
        r = para.add_run(value)
        _apply_font(r)


def _append_run(para, value: str):
    r = para.add_run(value.upper())
    _apply_font(r)


def _fill_fecha(doc):
    hoy = date.today()
    for p in doc.paragraphs:
        if "En Salamanca" in p.text and "XX" in p.text:
            runs = p.runs
            # runs: ['  En Salamanca', ' a ', 'XX', ' de', ' ', 'XXXXXXXXX', ' de 2026', '.          ']
            if len(runs) > 2:
                runs[2].text = str(hoy.day)
                _apply_font(runs[2])
            if len(runs) > 5:
                runs[5].text = _MESES[hoy.month - 1]
                _apply_font(runs[5])
            if len(runs) > 6:
                runs[6].text = f" de {hoy.year}"
                _apply_font(runs[6])
            break


def _generar_docx(familia, perro) -> bytes:
    doc = Document(os.path.abspath(_TEMPLATE_PATH))
    t0 = doc.tables[0]  # datos familia
    t1 = doc.tables[1]  # datos perro

    # ── Tabla 0: datos de la familia ────────────────────────────────────────
    # Fila 0 (fusionada): NOMBRE Y APELLIDOS — run[1]
    _set_run(t0.rows[0].cells[0].paragraphs[0], 1, f"{familia.nombre} {familia.apellidos}")

    # Fila 1: DNI (append) | CORREO (run[1])
    _append_run(t0.rows[1].cells[0].paragraphs[0], familia.dni or "")
    _set_run(t0.rows[1].cells[1].paragraphs[0], 1, familia.email or "")

    # Fila 2 (fusionada): DIRECCIÓN — run[1]
    _set_run(t0.rows[2].cells[0].paragraphs[0], 1, familia.direccion or "")

    # Fila 3: LOCALIDAD (run[1]) | PROVINCIA (run[1])
    _set_run(t0.rows[3].cells[0].paragraphs[0], 1, familia.municipio or "")
    _set_run(t0.rows[3].cells[1].paragraphs[0], 1, familia.provincia or "")

    # Fila 4: C.P (run[1]) | TELÉFONO (run[1])
    _set_run(t0.rows[4].cells[0].paragraphs[0], 1, familia.codigo_postal or "")
    _set_run(t0.rows[4].cells[1].paragraphs[0], 1, familia.telefono or "")

    # ── Tabla 1: datos del perro ─────────────────────────────────────────────
    # Fila 0 (fusionada): NOMBRE — run[1]
    _set_run(t1.rows[0].cells[0].paragraphs[0], 1, perro.nombre)

    # Fila 1: MICROCHIP (append, fusionada) | Nº PASAPORTE (append)
    _append_run(t1.rows[1].cells[0].paragraphs[0], perro.num_chip or "")
    _append_run(t1.rows[1].cells[2].paragraphs[0], perro.num_pasaporte or "")

    # Fila 2: RAZA (append) | SEXO (append) | F.NACIMIENTO (append)
    _append_run(t1.rows[2].cells[0].paragraphs[0], perro.raza.nombre if perro.raza else "")
    _append_run(t1.rows[2].cells[1].paragraphs[0], perro.sexo.value if perro.sexo else "")
    fecha_str = perro.fecha_nacimiento.strftime("%d/%m/%Y") if perro.fecha_nacimiento else ""
    _append_run(t1.rows[2].cells[2].paragraphs[0], fecha_str)

    # Fila 3: CAPA (append) | TAMAÑO (append) | ESTERILIZADO (append)
    _append_run(t1.rows[3].cells[0].paragraphs[0], perro.color or "")
    _append_run(t1.rows[3].cells[1].paragraphs[0], perro.tamano or "")
    _append_run(t1.rows[3].cells[2].paragraphs[0], "SÍ" if perro.esterilizado else "NO")

    # ── Párrafo fecha firma ──────────────────────────────────────────────────
    _fill_fecha(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_contrato_preadopcion(familia, perro) -> tuple[bytes | None, bytes]:
    """Devuelve (pdf_bytes_o_None, docx_bytes)."""
    from app.utils.pdf_utils import docx_a_pdf

    docx_bytes = _generar_docx(familia, perro)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        docx_path = tmp.name

    try:
        pdf_bytes = docx_a_pdf(docx_path)
    except Exception as e:
        logger.error("Error convirtiendo contrato preadopción a PDF: %s", e)
        pdf_bytes = None
    finally:
        os.unlink(docx_path)

    return pdf_bytes, docx_bytes
