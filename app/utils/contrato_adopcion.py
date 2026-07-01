import io
import logging
import os
import re
import tempfile
from datetime import date

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "contracts", "contrato_adopcion.docx")
_MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


_DATA_FONT = "Times New Roman"
_BASE_SIZE_PT = 10.0
_MIN_SIZE_PT = 6.0
_SIZE_STEP_PT = 0.5
_CELL_MARGIN_PT = 12.0  # margen izq+der de celda de Word por defecto, como colchón de seguridad

# Ancho de cada carácter en Times New Roman Negrita, como fracción del tamaño de fuente (em).
# Medido con la fuente real (timesbd.ttf) para no depender de tener la fuente instalada en producción.
_CHAR_WIDTH_EM = {
    " ": 0.25, "A": 0.722, "B": 0.667, "C": 0.722, "D": 0.722, "E": 0.667, "F": 0.611,
    "G": 0.778, "H": 0.778, "I": 0.389, "J": 0.5, "K": 0.778, "L": 0.667, "M": 0.944,
    "N": 0.722, "O": 0.778, "P": 0.611, "Q": 0.778, "R": 0.722, "S": 0.556, "T": 0.667,
    "U": 0.722, "V": 0.722, "W": 1.0, "X": 0.722, "Y": 0.722, "Z": 0.667,
    "a": 0.5, "b": 0.556, "c": 0.444, "d": 0.556, "e": 0.444, "f": 0.333, "g": 0.5,
    "h": 0.556, "i": 0.278, "j": 0.333, "k": 0.556, "l": 0.278, "m": 0.833, "n": 0.556,
    "o": 0.5, "p": 0.556, "q": 0.556, "r": 0.444, "s": 0.389, "t": 0.333, "u": 0.556,
    "v": 0.5, "w": 0.722, "x": 0.5, "y": 0.5, "z": 0.444,
    "Á": 0.722, "É": 0.667, "Í": 0.389, "Ó": 0.778, "Ú": 0.722, "Ñ": 0.722,
    "á": 0.5, "é": 0.444, "í": 0.278, "ó": 0.5, "ú": 0.556, "ñ": 0.556, "Ü": 0.722, "ü": 0.556,
    "0": 0.5, "1": 0.5, "2": 0.5, "3": 0.5, "4": 0.5, "5": 0.5, "6": 0.5, "7": 0.5, "8": 0.5, "9": 0.5,
    ".": 0.25, ",": 0.25, "@": 0.93, "/": 0.278, ":": 0.333, "-": 0.333, "_": 0.5,
    "º": 0.33, "ª": 0.3, "(": 0.333, ")": 0.333,
}
_DEFAULT_CHAR_WIDTH_EM = 0.6  # para caracteres no listados


def _text_width_pt(text: str, size_pt: float) -> float:
    if not text:
        return 0.0
    return sum(_CHAR_WIDTH_EM.get(c, _DEFAULT_CHAR_WIDTH_EM) for c in text) * size_pt


def _run_size_pt(run) -> float:
    return run.font.size.pt if run.font.size else _BASE_SIZE_PT


def _fit_font_size(cell, prefix_text: str, value_text: str, base_size_pt: float) -> float:
    """Calcula el tamaño de letra (>= _MIN_SIZE_PT, <= base_size_pt) para que value_text
    quepa en una línea de `cell` sin desplazar el resto del contenido de la fila/tabla."""
    if cell is None or not value_text:
        return base_size_pt
    cell_width_pt = (cell.width or 0) / 12700
    available_pt = cell_width_pt - _CELL_MARGIN_PT - _text_width_pt(prefix_text, base_size_pt)
    if available_pt <= 0:
        return _MIN_SIZE_PT
    size = base_size_pt
    while size > _MIN_SIZE_PT and _text_width_pt(value_text, size) > available_pt:
        size -= _SIZE_STEP_PT
    return max(size, _MIN_SIZE_PT)


def _apply_font(run, size_pt: float = _BASE_SIZE_PT):
    run.font.name = _DATA_FONT
    run.font.bold = True
    run.font.size = Pt(size_pt)


def _set_run(para, run_idx: int, value: str, upper: bool = True, cell=None):
    value = value.upper() if upper else value.lower()
    runs = para.runs
    prefix = "".join(r.text for r in runs[:run_idx])
    base_size = _run_size_pt(runs[run_idx]) if run_idx < len(runs) else (
        _run_size_pt(runs[-1]) if runs else _BASE_SIZE_PT
    )
    size = _fit_font_size(cell, prefix, value, base_size)
    if run_idx < len(runs):
        runs[run_idx].text = value
        _apply_font(runs[run_idx], size)
    else:
        r = para.add_run(value)
        _apply_font(r, size)


def _append_run(para, value: str, cell=None):
    value = value.upper()
    runs = para.runs
    base_size = _run_size_pt(runs[-1]) if runs else _BASE_SIZE_PT
    size = _fit_font_size(cell, para.text, value, base_size)
    r = para.add_run(value)
    _apply_font(r, size)


def _fill_tasa(doc, tasa):
    tasa_str = f"{tasa:.2f}" if tasa is not None else "0.00"
    for p in doc.paragraphs:
        if "cantidad de" in p.text:
            for r in p.runs:
                if "____" in r.text:
                    r.text = r.text.replace("____", tasa_str)
                    _apply_font(r)
                    break
            break


def _fill_fecha(doc):
    hoy = date.today()
    valores = [str(hoy.day), _MESES[hoy.month - 1], str(hoy.year)]
    for p in doc.paragraphs:
        if "En Salamanca" in p.text:
            placeholders = [r for r in p.runs if re.fullmatch(r"\s*_+\s*", r.text)]
            for r, valor in zip(placeholders, valores):
                leading = len(r.text) - len(r.text.lstrip())
                trailing = len(r.text) - len(r.text.rstrip())
                r.text = (" " * leading) + valor + (" " * trailing)
                _apply_font(r)
            break


def _generar_docx(familia, perro) -> bytes:
    doc = Document(os.path.abspath(_TEMPLATE_PATH))
    t0 = doc.tables[0]  # datos familia
    t1 = doc.tables[1]  # datos perro

    # ── Tabla 0: datos de la familia ────────────────────────────────────────
    # Fila 0 (celda fusionada): NOMBRE Y APELLIDOS — run[1] es el valor
    c = t0.rows[0].cells[0]
    _set_run(c.paragraphs[0], 1, f"{familia.nombre} {familia.apellidos}", cell=c)

    # Fila 1: DNI (1 run) | CORREO ELECTRÓNICO (run[1] es valor)
    c = t0.rows[1].cells[0]
    _append_run(c.paragraphs[0], familia.dni or "", cell=c)
    c = t0.rows[1].cells[1]
    _set_run(c.paragraphs[0], 1, familia.email or "", upper=False, cell=c)

    # Fila 2 (celda fusionada): DIRECCIÓN — 1 run, añadir valor
    c = t0.rows[2].cells[0]
    _append_run(c.paragraphs[0], familia.direccion or "", cell=c)

    # Fila 3: LOCALIDAD (run[1] es valor) | PROVINCIA (1 run)
    c = t0.rows[3].cells[0]
    _set_run(c.paragraphs[0], 1, familia.municipio or "", cell=c)
    c = t0.rows[3].cells[1]
    _append_run(c.paragraphs[0], familia.provincia or "", cell=c)

    # Fila 4: C.P (1 run) | TELÉFONO (1 run)
    c = t0.rows[4].cells[0]
    _append_run(c.paragraphs[0], familia.codigo_postal or "", cell=c)
    c = t0.rows[4].cells[1]
    _append_run(c.paragraphs[0], familia.telefono or "", cell=c)

    # ── Tabla 1: datos del perro ─────────────────────────────────────────────
    # Fila 0 (cols 0-1 fusionadas): NOMBRE — run[1] limpiar espacio, run[2] valor
    c = t1.rows[0].cells[0]
    _set_run(c.paragraphs[0], 1, "", cell=c)
    _set_run(c.paragraphs[0], 2, perro.nombre, cell=c)

    # Fila 1 (cols 0-1 fusionadas): MICROCHIP — run[1] limpiar, run[2] valor
    c = t1.rows[1].cells[0]
    _set_run(c.paragraphs[0], 1, "", cell=c)
    _set_run(c.paragraphs[0], 2, perro.num_chip or "", cell=c)

    # Fila 1 (cols 2-3 fusionadas): Nº PASAPORTE — run[2] limpiar espacio, run[3] valor
    c = t1.rows[1].cells[2]
    _set_run(c.paragraphs[0], 2, "", cell=c)
    _set_run(c.paragraphs[0], 3, perro.num_pasaporte or "", cell=c)

    # Fila 2 col 0: RAZA — run[2] es valor (bold)
    c = t1.rows[2].cells[0]
    _set_run(c.paragraphs[0], 2, perro.raza.nombre if perro.raza else "", cell=c)

    # Fila 2 col 1: SEXO — run[1] es valor
    c = t1.rows[2].cells[1]
    _set_run(c.paragraphs[0], 1, perro.sexo.value.upper() if perro.sexo else "", cell=c)

    # Fila 2 (cols 2-3 fusionadas): F.NACIMIENTO — run[1] limpiar espacios, run[2] valor
    c = t1.rows[2].cells[2]
    _set_run(c.paragraphs[0], 1, "", cell=c)
    fecha_str = perro.fecha_nacimiento.strftime("%d/%m/%Y") if perro.fecha_nacimiento else ""
    _set_run(c.paragraphs[0], 2, fecha_str, cell=c)

    # Fila 3 col 0: CAPA — run[1] es valor
    c = t1.rows[3].cells[0]
    _set_run(c.paragraphs[0], 1, perro.color or "", cell=c)

    # Fila 3 col 1: TAMAÑO — run[1] es valor
    c = t1.rows[3].cells[1]
    _set_run(c.paragraphs[0], 1, perro.tamano or "", cell=c)

    # Fila 3 (cols 2-3 fusionadas): ESTERILIZADO — run[1] es valor (reemplaza "PEND. ADOP")
    c = t1.rows[3].cells[2]
    _set_run(c.paragraphs[0], 1, "SÍ" if perro.esterilizado else "PEND. ADOP", cell=c)

    # ── Párrafo 21: tasa adopción ────────────────────────────────────────────
    _fill_tasa(doc, perro.tasa)

    # ── Párrafo fecha firma ──────────────────────────────────────────────────
    _fill_fecha(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_contrato_adopcion(familia, perro) -> tuple[bytes | None, bytes]:
    """Devuelve (pdf_bytes_o_None, docx_bytes)."""
    from app.utils.pdf_utils import docx_a_pdf

    docx_bytes = _generar_docx(familia, perro)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        docx_path = tmp.name

    try:
        pdf_bytes = docx_a_pdf(docx_path)
    except Exception as e:
        logger.error("Error convirtiendo contrato adopción a PDF: %s", e)
        pdf_bytes = None
    finally:
        os.unlink(docx_path)

    return pdf_bytes, docx_bytes
