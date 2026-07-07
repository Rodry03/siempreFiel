import io
import logging
import os
import tempfile
from datetime import date

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

logger = logging.getLogger(__name__)

_TURNO_LABELS = {"manana": "Mañana", "tarde": "Tarde"}


def _fmt_fecha(f) -> str:
    return f.strftime("%d/%m/%Y") if f else ""


def _calcular_edad(fecha_nacimiento) -> str:
    if not fecha_nacimiento:
        return ""
    hoy = date.today()
    anos = hoy.year - fecha_nacimiento.year - ((hoy.month, hoy.day) < (fecha_nacimiento.month, fecha_nacimiento.day))
    if anos == 0:
        meses = (hoy.year - fecha_nacimiento.year) * 12 + hoy.month - fecha_nacimiento.month
        if hoy.day < fecha_nacimiento.day:
            meses -= 1
        return f"{meses} mes{'es' if meses != 1 else ''}"
    return f"{anos} año{'s' if anos != 1 else ''}"


def _add_kv_table(doc, pares: list[tuple[str, str]]):
    filas = [(k, v) for k, v in pares if v]
    if not filas:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in filas:
        row = table.add_row()
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = v
    doc.add_paragraph()


def _add_data_table(doc, headers: list[str], filas: list[list[str]]):
    if not filas:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
    for fila in filas:
        row = table.add_row()
        for i, valor in enumerate(fila):
            row.cells[i].text = valor
    doc.add_paragraph()


def _add_foto(doc, foto_url: str):
    if not foto_url:
        return
    try:
        resp = requests.get(foto_url, timeout=10)
        resp.raise_for_status()
    except Exception as e:
        logger.warning("No se pudo descargar la foto del perro para la ficha: %s", e)
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(resp.content), width=Inches(2.5))
        doc.add_paragraph()
    except Exception as e:
        logger.warning("No se pudo insertar la foto del perro en la ficha: %s", e)


def _generar_docx(perro) -> bytes:
    doc = Document()

    titulo = perro.nombre
    if perro.nombre_nuevo:
        titulo = f"{perro.nombre} / {perro.nombre_nuevo}"
    doc.add_heading(f"Ficha del perro — {titulo}", level=1)

    _add_foto(doc, perro.foto_url)

    doc.add_heading("Datos básicos", level=2)
    _add_kv_table(doc, [
        ("Raza", perro.raza.nombre if perro.raza else ""),
        ("Sexo", perro.sexo.value if perro.sexo else ""),
        ("Fecha de nacimiento", _fmt_fecha(perro.fecha_nacimiento)),
        ("Edad", _calcular_edad(perro.fecha_nacimiento)),
        ("Tamaño", perro.tamano or ""),
        ("Color", perro.color or ""),
        ("Nº microchip", perro.num_chip or ""),
        ("Nº pasaporte", perro.num_pasaporte or ""),
        ("Esterilizado", "Sí" if perro.esterilizado else "No"),
        ("PPP", "Sí" if perro.ppp else "No"),
        ("Estado", perro.estado.value if perro.estado else ""),
        ("Fecha de entrada", _fmt_fecha(perro.fecha_entrada)),
        ("Fecha de adopción", _fmt_fecha(perro.fecha_adopcion)),
    ])

    if perro.notas and perro.notas.strip():
        doc.add_heading("Notas generales", level=2)
        doc.add_paragraph(perro.notas)
        doc.add_paragraph()

    ubicaciones = sorted(perro.ubicaciones, key=lambda u: u.fecha_inicio, reverse=True)
    if ubicaciones:
        doc.add_heading("Historial de ubicaciones", level=2)
        filas = []
        for u in ubicaciones:
            familia_o_contacto = ""
            if u.familia:
                familia_o_contacto = f"{u.familia.nombre} {u.familia.apellidos}"
            elif u.nombre_contacto:
                familia_o_contacto = u.nombre_contacto
            filas.append([
                u.tipo.value if u.tipo else "",
                _fmt_fecha(u.fecha_inicio),
                _fmt_fecha(u.fecha_fin) if u.fecha_fin else "Actual",
                familia_o_contacto,
            ])
        _add_data_table(doc, ["Tipo", "Desde", "Hasta", "Familia / Contacto"], filas)

    vacunas = sorted(perro.vacunas, key=lambda v: v.fecha_administracion, reverse=True)
    if vacunas:
        doc.add_heading("Vacunas", level=2)
        filas = [
            [v.tipo, _fmt_fecha(v.fecha_administracion), _fmt_fecha(v.fecha_proxima), v.veterinario or ""]
            for v in vacunas
        ]
        _add_data_table(doc, ["Tipo", "Administración", "Próxima", "Veterinario"], filas)

    if perro.pesos:
        doc.add_heading("Peso", level=2)
        filas = [[_fmt_fecha(p.fecha), f"{p.peso_kg} kg", p.notas or ""] for p in perro.pesos]
        _add_data_table(doc, ["Fecha", "Peso", "Notas"], filas)

    if perro.celos:
        doc.add_heading("Celos", level=2)
        filas = [
            [_fmt_fecha(c.fecha_inicio), _fmt_fecha(c.fecha_fin), c.notas or ""]
            for c in perro.celos
        ]
        _add_data_table(doc, ["Desde", "Hasta", "Notas"], filas)

    if perro.medicaciones:
        doc.add_heading("Medicación", level=2)
        filas = []
        for m in perro.medicaciones:
            turnos = (m.turno or "").split(",") if m.turno else []
            turno_str = ", ".join(_TURNO_LABELS.get(t, t) for t in turnos if t)
            filas.append([
                m.medicamento,
                m.dosis or "",
                m.frecuencia or "",
                turno_str,
                _fmt_fecha(m.fecha_inicio),
                _fmt_fecha(m.fecha_fin) if m.fecha_fin else "En curso",
            ])
        _add_data_table(doc, ["Medicamento", "Dosis", "Frecuencia", "Turno", "Desde", "Hasta"], filas)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_ficha_perro(perro) -> tuple[bytes | None, bytes]:
    """Devuelve (pdf_bytes_o_None, docx_bytes)."""
    from app.utils.pdf_utils import docx_a_pdf

    docx_bytes = _generar_docx(perro)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        docx_path = tmp.name

    try:
        pdf_bytes = docx_a_pdf(docx_path)
    except Exception as e:
        logger.error("Error convirtiendo ficha de perro a PDF: %s", e)
        pdf_bytes = None
    finally:
        os.unlink(docx_path)

    return pdf_bytes, docx_bytes
