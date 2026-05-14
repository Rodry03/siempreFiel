import io
import logging
import os
import subprocess
import tempfile

logger = logging.getLogger(__name__)
from datetime import date
from docx import Document as DocxDocument
import cloudinary
import cloudinary.uploader
from fastapi import APIRouter, Depends, File, Form, Query, Request, UploadFile
from app.auth import get_current_user, require_not_veterano, flash
from fastapi.responses import RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from typing import List, Optional
from app.database import get_db
from app.models import Voluntario, PerfilVoluntario, EstadoContrato, GrupoTarea, MiembroGrupoTarea, PeriodoApoyo
from app.templates_config import templates

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "contracts", "contrato_voluntario.docx")
MESES = ["enero","febrero","marzo","abril","mayo","junio","julio","agosto","septiembre","octubre","noviembre","diciembre"]


def _set_cell_text(cell, text):
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _set_paragraph_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _docx_a_pdf(docx_path: str) -> bytes | None:
    # Intenta con Word COM (Windows) — inicializa COM para contexto multi-hilo
    try:
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            with tempfile.TemporaryDirectory() as tmpdir:
                pdf_path = os.path.join(tmpdir, "contrato.pdf")
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
                doc.Close(False)
                with open(pdf_path, "rb") as f:
                    return f.read()
        finally:
            if word:
                word.Quit()
            pythoncom.CoUninitialize()
    except Exception as e:
        logger.error("Word COM falló: %s: %s", type(e).__name__, e)

    # Fallback: LibreOffice (Linux/Render)
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [
                    "/usr/bin/soffice",
                    "--headless",
                    "-env:UserInstallation=file:///tmp/libreoffice_profile",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                capture_output=True, timeout=60,
            )
            logger.warning("soffice returncode: %s", result.returncode)
            logger.warning("soffice stdout: %s", result.stdout.decode(errors="replace"))
            logger.warning("soffice stderr: %s", result.stderr.decode(errors="replace"))
            if result.returncode != 0:
                return None
            base = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_path = os.path.join(tmpdir, base + ".pdf")
            with open(pdf_path, "rb") as f:
                return f.read()
    except FileNotFoundError:
        logger.warning("soffice no encontrado en el sistema")
        return None

router = APIRouter(prefix="/voluntarios", dependencies=[Depends(get_current_user), Depends(require_not_veterano)])


@router.get("/debug/soffice")
def debug_soffice():
    result = subprocess.run(
        ["find", "/", "-name", "soffice", "-type", "f"],
        capture_output=True, timeout=30,
    )
    return {
        "stdout": result.stdout.decode(errors="replace"),
        "stderr": result.stderr.decode(errors="replace"),
    }

PERFIL_LABELS = {
    "directiva": "Directiva",
    "apoyo_en_junta": "Apoyo en Junta",
    "veterano": "Veterano",
    "voluntario": "Voluntario",
    "guagua": "Guagua",
    "eventos": "Eventos",
    "colaboradores": "Colaboradores",
}
PERFIL_COLORS = {
    "directiva": "danger",
    "apoyo_en_junta": "dark",
    "veterano": "warning",
    "voluntario": "success",
    "guagua": "primary",
    "eventos": "info",
    "colaboradores": "secondary",
}


@router.get("/")
def listar_voluntarios(
    request: Request,
    perfil: List[str] = Query(default=[]),
    bajas: bool = False,
    sort: str = "nombre",
    order: str = "asc",
    db: Session = Depends(get_db),
):
    query = db.query(Voluntario)
    if not bajas:
        query = query.filter(Voluntario.activo == True)
    valores_validos = {p.value for p in PerfilVoluntario}
    perfiles_filtro = [p for p in perfil if p in valores_validos]
    if perfiles_filtro:
        query = query.filter(Voluntario.perfil.in_(perfiles_filtro))

    _cols = {
        "nombre": Voluntario.nombre,
        "apellido": Voluntario.apellido,
        "email": Voluntario.email,
        "perfil": Voluntario.perfil,
        "fecha_alta": Voluntario.fecha_alta,
    }
    col = _cols.get(sort, Voluntario.nombre)
    query = query.order_by(col.desc() if order == "desc" else col.asc())

    perfil_qs = "&".join(f"perfil={p}" for p in perfiles_filtro)

    from app.routers.turnos import calcular_saldo, PERFILES_SIN_TURNOS
    voluntarios = query.all()
    saldos = {
        v.id: calcular_saldo(v)
        for v in voluntarios
        if v.perfil not in PERFILES_SIN_TURNOS
    }

    return templates.TemplateResponse(request, "voluntarios/list.html", {
        "voluntarios": voluntarios,
        "saldos": saldos,
        "perfil_filtro": perfiles_filtro,
        "perfil_qs": perfil_qs,
        "perfiles": [p.value for p in PerfilVoluntario],
        "perfil_labels": PERFIL_LABELS,
        "perfil_colors": PERFIL_COLORS,
        "bajas": bajas,
        "sort": sort,
        "order": order,
    })


CONTRATO_LABELS = {
    "pendiente": "Pendiente",
    "enviado": "Enviado",
    "firmado": "Firmado",
}
CONTRATO_COLORS = {
    "pendiente": "secondary",
    "enviado": "warning",
    "firmado": "success",
}


def _contexto_form(extra={}):
    return {
        "perfiles": [p.value for p in PerfilVoluntario],
        "perfil_labels": PERFIL_LABELS,
        "contratos": [e.value for e in EstadoContrato],
        "contrato_labels": CONTRATO_LABELS,
        "hoy": date.today().isoformat(),
        **extra,
    }


@router.get("/nuevo")
def form_nuevo_voluntario(
    request: Request,
    nombre: str = "",
    apellido: str = "",
    email: str = "",
    telefono: str = "",
):
    prefill = {"nombre": nombre, "apellido": apellido, "email": email, "telefono": telefono}
    return templates.TemplateResponse(request, "voluntarios/form.html",
        _contexto_form({"voluntario": None, "prefill": prefill}))


@router.post("/nuevo")
def crear_voluntario(
    request: Request,
    nombre: str = Form(...),
    apellido: str = Form(...),
    dni: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    perfil: str = Form(...),
    fecha_alta: date = Form(...),
    activo: Optional[str] = Form(None),
    ppp: Optional[str] = Form(None),
    telefono: str = Form(...),
    direccion: Optional[str] = Form(None),
    ciudad: Optional[str] = Form(None),
    provincia: Optional[str] = Form(None),
    codigo_postal: Optional[str] = Form(None),
    fecha_contrato: Optional[date] = Form(None),
    contrato_estado: Optional[str] = Form(None),
    teaming: Optional[str] = Form(None),
    notas: Optional[str] = Form(None),
    fecha_veterano: Optional[date] = Form(None),
    fecha_fin_veterano: Optional[date] = Form(None),
    db: Session = Depends(get_db),
):
    email_final = email or f"{nombre}{apellido}@siemprefiel.com"
    if fecha_fin_veterano:
        perfil = PerfilVoluntario.voluntario.value
    voluntario = Voluntario(
        nombre=nombre,
        apellido=apellido,
        dni=dni or None,
        email=email_final,
        perfil=PerfilVoluntario(perfil),
        fecha_alta=fecha_alta,
        activo=activo == "on",
        ppp=ppp == "on",
        telefono=telefono,
        direccion=direccion or None,
        ciudad=ciudad or None,
        provincia=provincia or None,
        codigo_postal=codigo_postal or None,
        fecha_contrato=fecha_contrato,
        contrato_estado=EstadoContrato(contrato_estado) if contrato_estado else None,
        teaming=teaming == "on",
        notas=notas or None,
        fecha_veterano=fecha_veterano,
        fecha_fin_veterano=fecha_fin_veterano,
    )
    db.add(voluntario)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        return templates.TemplateResponse(request, "voluntarios/form.html",
            _contexto_form({"voluntario": voluntario, "error": "El email o DNI ya está registrado."}))
    flash(request, f"Voluntario {voluntario.nombre} {voluntario.apellido} creado correctamente.")
    return RedirectResponse("/voluntarios/", status_code=303)


@router.get("/{voluntario_id}/editar")
def form_editar_voluntario(request: Request, voluntario_id: int, db: Session = Depends(get_db)):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if not voluntario:
        return RedirectResponse("/voluntarios/", status_code=303)
    return templates.TemplateResponse(request, "voluntarios/form.html",
        _contexto_form({"voluntario": voluntario}))


@router.post("/{voluntario_id}/editar")
def editar_voluntario(
    request: Request,
    voluntario_id: int,
    nombre: str = Form(...),
    apellido: str = Form(...),
    dni: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    perfil: str = Form(...),
    fecha_alta: date = Form(...),
    activo: Optional[str] = Form(None),
    ppp: Optional[str] = Form(None),
    telefono: str = Form(...),
    direccion: Optional[str] = Form(None),
    ciudad: Optional[str] = Form(None),
    provincia: Optional[str] = Form(None),
    codigo_postal: Optional[str] = Form(None),
    fecha_contrato: Optional[date] = Form(None),
    contrato_estado: Optional[str] = Form(None),
    teaming: Optional[str] = Form(None),
    notas: Optional[str] = Form(None),
    fecha_veterano: Optional[date] = Form(None),
    fecha_fin_veterano: Optional[date] = Form(None),
    db: Session = Depends(get_db),
):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if not voluntario:
        return RedirectResponse("/voluntarios/", status_code=303)
    email_final = email or f"{nombre}{apellido}@siemprefiel.com"
    voluntario.nombre = nombre
    voluntario.apellido = apellido
    voluntario.dni = dni or None
    voluntario.email = email_final
    voluntario.fecha_alta = fecha_alta
    voluntario.activo = activo == "on"
    voluntario.ppp = ppp == "on"
    voluntario.telefono = telefono
    voluntario.direccion = direccion or None
    voluntario.ciudad = ciudad or None
    voluntario.provincia = provincia or None
    voluntario.codigo_postal = codigo_postal or None
    voluntario.fecha_contrato = fecha_contrato
    voluntario.contrato_estado = EstadoContrato(contrato_estado) if contrato_estado else None
    voluntario.teaming = teaming == "on"
    voluntario.notas = notas or None
    if fecha_fin_veterano:
        perfil = PerfilVoluntario.voluntario.value
    voluntario.perfil = PerfilVoluntario(perfil)
    voluntario.fecha_veterano = fecha_veterano
    voluntario.fecha_fin_veterano = fecha_fin_veterano
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        return templates.TemplateResponse(request, "voluntarios/form.html",
            _contexto_form({"voluntario": voluntario, "error": "El email o DNI ya está registrado."}))
    flash(request, "Cambios guardados.")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.get("/{voluntario_id}/contrato")
def generar_contrato(voluntario_id: int, db: Session = Depends(get_db)):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if not voluntario:
        return RedirectResponse("/voluntarios/", status_code=303)

    doc = DocxDocument(TEMPLATE_PATH)
    tabla = doc.tables[0]

    hoy = date.today()
    fecha_str = f"  En Salamanca a {hoy.day} de {MESES[hoy.month - 1]} del {hoy.year}.          "

    _set_cell_text(tabla.rows[0].cells[0], f"NOMBRE Y APELLIDOS: {voluntario.nombre} {voluntario.apellido}")
    _set_cell_text(tabla.rows[1].cells[0], f"DNI: {voluntario.dni or ''}")
    _set_cell_text(tabla.rows[1].cells[1], f"CORREO ELECTRÓNICO: {voluntario.email or ''}")
    _set_cell_text(tabla.rows[2].cells[0], f"DIRECCIÓN: {voluntario.direccion or ''}")
    _set_cell_text(tabla.rows[3].cells[1], f"PROVINCIA: {voluntario.provincia or 'Salamanca'}")
    _set_cell_text(tabla.rows[4].cells[0], f"C.P: {voluntario.codigo_postal or '37004'}")
    _set_cell_text(tabla.rows[4].cells[1], f"TELÉFONO: {voluntario.telefono or ''}")

    for para in doc.paragraphs:
        if "En Salamanca a" in para.text:
            _set_paragraph_text(para, fecha_str)
            break

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    nombre = f"{voluntario.nombre}_{voluntario.apellido}"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        docx_path = tmp.name

    try:
        pdf_bytes = _docx_a_pdf(docx_path)
    finally:
        os.unlink(docx_path)

    if pdf_bytes:
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="Contrato_{nombre}.pdf"'},
        )

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Contrato_{nombre}.docx"'},
    )


_MAX_CONTRATO_BYTES = 10 * 1024 * 1024  # 10 MB


def _subir_contrato_firmado(file: UploadFile, voluntario_id: int) -> tuple:
    cloudinary.config(
        cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
        api_key=os.environ.get("CLOUDINARY_API_KEY"),
        api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
    )
    contents = file.file.read(_MAX_CONTRATO_BYTES + 1)
    if len(contents) > _MAX_CONTRATO_BYTES:
        raise ValueError("El archivo supera el tamaño máximo permitido (10 MB).")
    result = cloudinary.uploader.upload(
        contents,
        resource_type="raw",
        folder="protectora/contratos",
        public_id=f"contrato_voluntario_{voluntario_id}",
        overwrite=True,
    )
    return result["secure_url"], file.filename


@router.get("/contratos-firmados")
def contratos_firmados(request: Request, db: Session = Depends(get_db)):
    voluntarios = db.query(Voluntario).filter(Voluntario.activo == True).order_by(Voluntario.nombre.asc()).all()
    return templates.TemplateResponse(request, "voluntarios/contratos_firmados.html", {
        "voluntarios": voluntarios,
        "perfil_labels": PERFIL_LABELS,
        "perfil_colors": PERFIL_COLORS,
        "contrato_labels": CONTRATO_LABELS,
        "contrato_colors": CONTRATO_COLORS,
    })


@router.post("/{voluntario_id}/contrato-firmado")
def subir_contrato_firmado(
    request: Request,
    voluntario_id: int,
    archivo: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if not voluntario:
        return RedirectResponse("/voluntarios/", status_code=303)
    try:
        url, nombre = _subir_contrato_firmado(archivo, voluntario_id)
        voluntario.contrato_firmado_url = url
        voluntario.contrato_firmado_fecha = date.today()
        voluntario.contrato_firmado_nombre = nombre
        db.commit()
        flash(request, "Contrato firmado subido correctamente.")
    except Exception as e:
        logger.error("Error subiendo contrato firmado: %s", e)
        flash(request, "Error al subir el contrato.", "danger")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/contrato-firmado/eliminar")
def eliminar_contrato_firmado(
    request: Request,
    voluntario_id: int,
    db: Session = Depends(get_db),
):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if voluntario and voluntario.contrato_firmado_url:
        try:
            cloudinary.config(
                cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
                api_key=os.environ.get("CLOUDINARY_API_KEY"),
                api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
            )
            cloudinary.uploader.destroy(
                f"protectora/contratos/contrato_voluntario_{voluntario_id}",
                resource_type="raw",
            )
        except Exception as e:
            logger.warning("Error eliminando contrato de Cloudinary: %s", e)
        voluntario.contrato_firmado_url = None
        voluntario.contrato_firmado_fecha = None
        voluntario.contrato_firmado_nombre = None
        db.commit()
        flash(request, "Contrato eliminado.")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/dar-de-baja")
def dar_de_baja(request: Request, voluntario_id: int, db: Session = Depends(get_db)):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if voluntario:
        voluntario.activo = False
        db.query(GrupoTarea).filter(GrupoTarea.capitan_id == voluntario_id).update({"capitan_id": None})
        db.query(MiembroGrupoTarea).filter(MiembroGrupoTarea.voluntario_id == voluntario_id).delete()
        db.commit()
        flash(request, f"{voluntario.nombre} {voluntario.apellido} dado/a de baja.", "warning")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/reactivar")
def reactivar(request: Request, voluntario_id: int, db: Session = Depends(get_db)):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if voluntario:
        voluntario.activo = True
        db.commit()
        flash(request, f"{voluntario.nombre} {voluntario.apellido} reactivado/a.")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/cambiar-perfil")
def cambiar_perfil(voluntario_id: int, perfil: str = Form(...), db: Session = Depends(get_db)):
    voluntario = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if voluntario:
        try:
            voluntario.perfil = PerfilVoluntario(perfil)
            db.commit()
        except ValueError:
            pass
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/apoyo/nuevo")
def nuevo_periodo_apoyo(
    request: Request,
    voluntario_id: int,
    fecha_inicio: date = Form(...),
    fecha_fin: Optional[date] = Form(None),
    db: Session = Depends(get_db),
):
    db.add(PeriodoApoyo(voluntario_id=voluntario_id, fecha_inicio=fecha_inicio, fecha_fin=fecha_fin))
    db.commit()
    flash(request, "Periodo de apoyo registrado.")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/apoyo/{periodo_id}/cerrar")
def cerrar_periodo_apoyo(
    request: Request,
    voluntario_id: int,
    periodo_id: int,
    fecha_fin: date = Form(...),
    db: Session = Depends(get_db),
):
    p = db.query(PeriodoApoyo).filter(PeriodoApoyo.id == periodo_id).first()
    if p:
        p.fecha_fin = fecha_fin
        db.commit()
        flash(request, "Periodo de apoyo cerrado.")
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/apoyo/{periodo_id}/eliminar")
def eliminar_periodo_apoyo(
    request: Request,
    voluntario_id: int,
    periodo_id: int,
    db: Session = Depends(get_db),
):
    p = db.query(PeriodoApoyo).filter(PeriodoApoyo.id == periodo_id).first()
    if p:
        db.delete(p)
        db.commit()
    return RedirectResponse(f"/voluntarios/{voluntario_id}", status_code=303)


@router.post("/{voluntario_id}/eliminar")
def eliminar_voluntario(voluntario_id: int, request: Request, db: Session = Depends(get_db)):
    from app.models import EjecucionGrupoTarea, Usuario, NotaGestion, SaldoMensual
    v = db.query(Voluntario).filter(Voluntario.id == voluntario_id).first()
    if not v:
        return RedirectResponse("/voluntarios/", status_code=303)
    nombre = f"{v.nombre} {v.apellido}"
    db.query(GrupoTarea).filter(GrupoTarea.capitan_id == voluntario_id).update({"capitan_id": None})
    db.query(EjecucionGrupoTarea).filter(EjecucionGrupoTarea.ejecutor_id == voluntario_id).update({"ejecutor_id": None})
    db.query(Usuario).filter(Usuario.voluntario_id == voluntario_id).update({"voluntario_id": None})
    db.query(NotaGestion).filter(NotaGestion.encargado_id == voluntario_id).update({"encargado_id": None})
    db.query(MiembroGrupoTarea).filter(MiembroGrupoTarea.voluntario_id == voluntario_id).delete()
    db.query(SaldoMensual).filter(SaldoMensual.voluntario_id == voluntario_id).delete()
    db.delete(v)
    db.commit()
    flash(request, f"{nombre} eliminado/a.", "success")
    return RedirectResponse("/voluntarios/", status_code=303)
